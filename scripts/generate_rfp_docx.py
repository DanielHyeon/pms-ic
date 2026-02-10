#!/usr/bin/env python3
"""
RFP Word 문서 생성기
- DB에서 프로젝트별 WBS/요구사항 데이터를 조회하여 RFP .docx 파일을 생성한다.
- 생성된 파일은 RFP 업로드 파이프라인 테스트용으로 사용한다.
"""

import os
import psycopg2
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# DB 연결 정보 (docker-compose 기준)
DB_CONFIG = {
    "host": "localhost",
    "port": 5433,
    "database": "pms_db",
    "user": "pms_user",
    "password": "pms_password",
}

# 출력 디렉토리
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "test-data", "rfp-docs")


def get_connection():
    """PostgreSQL 연결"""
    return psycopg2.connect(**DB_CONFIG)


def fetch_projects(conn):
    """전체 프로젝트 목록 조회"""
    with conn.cursor() as cur:
        cur.execute("""
            SELECT id, name, description, status, start_date, end_date, budget
            FROM project.projects
            ORDER BY name
        """)
        columns = [desc[0] for desc in cur.description]
        return [dict(zip(columns, row)) for row in cur.fetchall()]


def fetch_phases(conn, project_id):
    """프로젝트의 페이즈 조회 (부모-자식 계층 포함)"""
    with conn.cursor() as cur:
        cur.execute("""
            SELECT id, parent_id, name, description, start_date, end_date,
                   progress, order_num, status
            FROM project.phases
            WHERE project_id = %s
            ORDER BY order_num
        """, (project_id,))
        columns = [desc[0] for desc in cur.description]
        return [dict(zip(columns, row)) for row in cur.fetchall()]


def fetch_wbs_groups(conn, project_id):
    """프로젝트의 WBS 그룹 조회"""
    with conn.cursor() as cur:
        cur.execute("""
            SELECT g.id, g.phase_id, g.code, g.name, g.description, g.status,
                   g.planned_start_date, g.planned_end_date, g.progress, g.order_num
            FROM project.wbs_groups g
            WHERE g.project_id = %s
            ORDER BY g.phase_id, g.order_num
        """, (project_id,))
        columns = [desc[0] for desc in cur.description]
        return [dict(zip(columns, row)) for row in cur.fetchall()]


def fetch_wbs_items(conn, project_id):
    """프로젝트의 WBS 아이템 조회"""
    with conn.cursor() as cur:
        cur.execute("""
            SELECT i.id, i.group_id, i.code, i.name, i.description, i.status,
                   i.planned_start_date, i.planned_end_date, i.estimated_hours, i.order_num
            FROM project.wbs_items i
            JOIN project.wbs_groups g ON i.group_id = g.id
            WHERE g.project_id = %s
            ORDER BY g.phase_id, g.order_num, i.order_num
        """, (project_id,))
        columns = [desc[0] for desc in cur.description]
        return [dict(zip(columns, row)) for row in cur.fetchall()]


def fetch_wbs_tasks(conn, project_id):
    """프로젝트의 WBS 태스크 조회"""
    with conn.cursor() as cur:
        cur.execute("""
            SELECT t.id, t.item_id, t.code, t.name, t.description, t.status,
                   t.planned_start_date, t.planned_end_date, t.estimated_hours, t.order_num
            FROM project.wbs_tasks t
            JOIN project.wbs_items i ON t.item_id = i.id
            JOIN project.wbs_groups g ON i.group_id = g.id
            WHERE g.project_id = %s
            ORDER BY g.phase_id, g.order_num, i.order_num, t.order_num
        """, (project_id,))
        columns = [desc[0] for desc in cur.description]
        return [dict(zip(columns, row)) for row in cur.fetchall()]


def fetch_requirements(conn, project_id):
    """프로젝트의 요구사항 조회"""
    with conn.cursor() as cur:
        cur.execute("""
            SELECT id, requirement_code, title, description, category,
                   priority, status, acceptance_criteria
            FROM project.requirements
            WHERE project_id = %s
            ORDER BY category, priority
        """, (project_id,))
        columns = [desc[0] for desc in cur.description]
        return [dict(zip(columns, row)) for row in cur.fetchall()]


def fetch_risks(conn, project_id):
    """프로젝트의 리스크 조회"""
    with conn.cursor() as cur:
        cur.execute("""
            SELECT id, title, description, category, impact, probability,
                   score, status, mitigation_plan
            FROM project.risks
            WHERE project_id = %s
            ORDER BY score DESC NULLS LAST, impact DESC NULLS LAST
        """, (project_id,))
        columns = [desc[0] for desc in cur.description]
        return [dict(zip(columns, row)) for row in cur.fetchall()]


# ── 문서 스타일 설정 ─────────────────────────────────────────

def set_document_styles(doc):
    """문서 기본 스타일 설정 (한글 폰트 포함)"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10)
    # 한글 폰트 설정
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # 제목 스타일
    for i in range(1, 5):
        heading = doc.styles[f"Heading {i}"]
        heading.font.name = "맑은 고딕"
        heading.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        heading.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)


def add_cover_page(doc, project):
    """표지 페이지 생성"""
    # 상단 여백
    for _ in range(6):
        doc.add_paragraph()

    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("제안 요청서 (RFP)")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    doc.add_paragraph()

    # 프로젝트명
    proj_name = doc.add_paragraph()
    proj_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = proj_name.add_run(project["name"])
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    # 메타 정보
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"문서 버전: v1.0\n작성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # 발주처 정보
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("발주처: 인슈어테크(주) AI사업부\n담당자: 프로젝트관리팀")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 페이지 구분
    doc.add_page_break()


def add_table_of_contents(doc):
    """목차 페이지"""
    doc.add_heading("목 차", level=1)
    doc.add_paragraph()

    toc_items = [
        "1. 사업 개요",
        "   1.1 사업 목적",
        "   1.2 사업 범위",
        "   1.3 추진 일정",
        "   1.4 예산",
        "2. 현행 업무 분석",
        "3. 요구사항 정의",
        "   3.1 기능 요구사항",
        "   3.2 비기능 요구사항",
        "   3.3 보안 요구사항",
        "   3.4 연동 요구사항",
        "4. WBS (작업분류체계)",
        "   4.1 페이즈별 WBS 구조",
        "   4.2 WBS 상세 내역",
        "5. 리스크 관리",
        "6. 제안 요구사항",
        "   6.1 제안서 구성",
        "   6.2 평가 기준",
        "7. 계약 조건",
    ]

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


def format_date(d):
    """날짜를 한글 형식으로 변환"""
    if d is None:
        return "-"
    return d.strftime("%Y.%m.%d")


def format_budget(amount):
    """예산을 한글 형식으로 변환 (억 단위)"""
    if amount is None:
        return "-"
    억 = int(amount) // 100_000_000
    만 = (int(amount) % 100_000_000) // 10_000
    if 억 > 0 and 만 > 0:
        return f"{억}억 {만:,}만원"
    elif 억 > 0:
        return f"{억}억원"
    else:
        return f"{만:,}만원"


def add_styled_table(doc, headers, rows, col_widths=None):
    """스타일이 적용된 테이블 생성"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 헤더 행 (진한 파란색 배경)
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        # 배경색 설정
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "1A365D",
        })
        shading.append(shading_elem)
        # 폰트 설정
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)

    # 데이터 행
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(value) if value is not None else "-"
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            # 짝수 행 배경색 (연한 회색)
            if row_idx % 2 == 0:
                shading = cell._element.get_or_add_tcPr()
                shading_elem = shading.makeelement(qn("w:shd"), {
                    qn("w:val"): "clear",
                    qn("w:color"): "auto",
                    qn("w:fill"): "F0F4FA",
                })
                shading.append(shading_elem)

    # 컬럼 너비 설정
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    return table


# ── 섹션별 내용 작성 ─────────────────────────────────────────

def add_section_1(doc, project):
    """1. 사업 개요"""
    doc.add_heading("1. 사업 개요", level=1)

    # 1.1 사업 목적
    doc.add_heading("1.1 사업 목적", level=2)
    doc.add_paragraph(project["description"] or "해당 프로젝트의 목적을 기술합니다.")
    doc.add_paragraph()

    # 주요 목표 (프로젝트 설명에서 파생)
    doc.add_heading("사업 목표", level=3)
    goals = [
        "업무 효율성 향상을 통한 처리 시간 단축",
        "AI/자동화 기술 도입으로 정확도 개선",
        "사용자 경험 개선 및 고객 만족도 향상",
        "시스템 안정성 및 보안 강화",
    ]
    for goal in goals:
        doc.add_paragraph(goal, style="List Bullet")

    doc.add_paragraph()

    # 1.2 사업 범위
    doc.add_heading("1.2 사업 범위", level=2)

    add_styled_table(doc,
        ["항목", "내용"],
        [
            ["사업명", project["name"]],
            ["사업 기간", f'{format_date(project["start_date"])} ~ {format_date(project["end_date"])}'],
            ["사업 예산", format_budget(project["budget"])],
            ["현재 상태", project["status"]],
            ["발주 기관", "인슈어테크(주)"],
        ],
        col_widths=[4, 12],
    )
    doc.add_paragraph()

    # 1.3 추진 일정
    doc.add_heading("1.3 추진 일정", level=2)
    doc.add_paragraph(
        "본 사업은 아래와 같은 일정으로 추진되며, "
        "상세 일정은 WBS(작업분류체계)에서 확인할 수 있습니다."
    )

    # 1.4 예산
    doc.add_heading("1.4 예산", level=2)
    doc.add_paragraph(
        f"본 사업의 총 예산은 {format_budget(project['budget'])}이며, "
        "세부 항목은 협의를 통해 조정될 수 있습니다."
    )
    doc.add_paragraph()


def add_section_2(doc, phases):
    """2. 현행 업무 분석"""
    doc.add_heading("2. 현행 업무 분석", level=1)
    doc.add_paragraph(
        "본 사업의 대상 업무를 분석한 결과, 다음과 같은 개선이 필요한 것으로 파악되었습니다."
    )
    doc.add_paragraph()

    # 페이즈를 기반으로 현행 업무 설명 생성
    # 최상위 페이즈만 표시
    top_phases = [p for p in phases if p["parent_id"] is None]
    rows = []
    for p in top_phases:
        rows.append([
            p["name"],
            p["description"] or "-",
            format_date(p["start_date"]),
            format_date(p["end_date"]),
            f'{p["progress"]}%' if p["progress"] is not None else "0%",
        ])

    if rows:
        add_styled_table(doc,
            ["단계", "설명", "시작일", "종료일", "진행률"],
            rows,
            col_widths=[3.5, 7, 2.5, 2.5, 1.5],
        )
    doc.add_paragraph()


def add_section_3(doc, requirements):
    """3. 요구사항 정의"""
    doc.add_heading("3. 요구사항 정의", level=1)
    doc.add_paragraph(
        "본 사업에서 요구하는 기능적/비기능적 요구사항은 다음과 같습니다. "
        "제안사는 아래 요구사항을 충족하는 방안을 제시하여야 합니다."
    )
    doc.add_paragraph()

    # 카테고리별 분류
    categories = {
        "FUNCTIONAL": ("3.1 기능 요구사항", "기능"),
        "NON_FUNCTIONAL": ("3.2 비기능 요구사항", "비기능"),
        "SECURITY": ("3.3 보안 요구사항", "보안"),
        "INTEGRATION": ("3.4 연동 요구사항", "연동"),
        "AI": ("3.5 AI/지능화 요구사항", "AI"),
    }

    for cat_key, (heading, _label) in categories.items():
        cat_reqs = [r for r in requirements if r["category"] == cat_key]
        if not cat_reqs:
            continue

        doc.add_heading(heading, level=2)

        for req in cat_reqs:
            # 요구사항 코드 + 제목
            code = req["requirement_code"] or req["id"]
            doc.add_heading(f"[{code}] {req['title']}", level=3)

            # 설명
            if req["description"]:
                doc.add_paragraph(req["description"])

            # 메타 정보 테이블
            add_styled_table(doc,
                ["항목", "내용"],
                [
                    ["우선순위", req["priority"] or "-"],
                    ["상태", req["status"] or "-"],
                    ["수용 기준", req["acceptance_criteria"] or "협의 필요"],
                ],
                col_widths=[4, 12],
            )
            doc.add_paragraph()

    # 요구사항이 없는 경우
    if not requirements:
        doc.add_paragraph("(요구사항은 사업 진행 과정에서 상세 정의 예정)")
    doc.add_paragraph()


def add_section_4(doc, phases, groups, items, tasks):
    """4. WBS (작업분류체계)"""
    doc.add_heading("4. WBS (작업분류체계)", level=1)
    doc.add_paragraph(
        "본 사업의 작업분류체계(WBS)는 다음과 같이 구성됩니다. "
        "각 단계별 세부 작업 항목과 일정, 예상 공수를 포함합니다."
    )
    doc.add_paragraph()

    # 4.1 페이즈별 WBS 구조
    doc.add_heading("4.1 페이즈별 WBS 구조", level=2)

    # 최상위 페이즈만 먼저 보여줌
    top_phases = [p for p in phases if p["parent_id"] is None]

    # 그룹을 phase_id로 인덱싱
    groups_by_phase = {}
    for g in groups:
        groups_by_phase.setdefault(g["phase_id"], []).append(g)

    # 아이템을 group_id로 인덱싱
    items_by_group = {}
    for item in items:
        items_by_group.setdefault(item["group_id"], []).append(item)

    # 태스크를 item_id로 인덱싱
    tasks_by_item = {}
    for task in tasks:
        tasks_by_item.setdefault(task["item_id"], []).append(task)

    # 하위 페이즈를 parent_id로 인덱싱
    child_phases = {}
    for p in phases:
        if p["parent_id"]:
            child_phases.setdefault(p["parent_id"], []).append(p)

    for phase in top_phases:
        doc.add_heading(f"■ {phase['name']}", level=3)
        if phase["description"]:
            doc.add_paragraph(phase["description"])

        # 기간 정보
        doc.add_paragraph(
            f"기간: {format_date(phase['start_date'])} ~ {format_date(phase['end_date'])}  |  "
            f"진행률: {phase['progress'] or 0}%"
        )

        # 해당 페이즈의 WBS 그룹
        phase_groups = groups_by_phase.get(phase["id"], [])

        # 하위 페이즈가 있으면 그것도 포함
        sub_phases = child_phases.get(phase["id"], [])
        for sp in sub_phases:
            sp_groups = groups_by_phase.get(sp["id"], [])
            phase_groups.extend(sp_groups)

        if phase_groups:
            rows = []
            for g in phase_groups:
                g_items = items_by_group.get(g["id"], [])
                item_names = ", ".join([it["name"] for it in g_items[:3]])
                if len(g_items) > 3:
                    item_names += f" 외 {len(g_items) - 3}건"
                rows.append([
                    g["code"] or "-",
                    g["name"],
                    g["description"] or "-",
                    g["status"] or "-",
                    format_date(g["planned_start_date"]),
                    format_date(g["planned_end_date"]),
                ])

            add_styled_table(doc,
                ["코드", "WBS 그룹", "설명", "상태", "시작일", "종료일"],
                rows,
                col_widths=[1.5, 3, 5, 2, 2.5, 2.5],
            )
        doc.add_paragraph()

    # 4.2 WBS 상세 내역
    doc.add_heading("4.2 WBS 상세 내역", level=2)
    doc.add_paragraph(
        "아래는 WBS 그룹별 상세 작업 항목(WBS Item)과 세부 태스크 목록입니다."
    )
    doc.add_paragraph()

    for group in groups:
        g_items = items_by_group.get(group["id"], [])
        if not g_items:
            continue

        doc.add_heading(f"▶ [{group['code']}] {group['name']}", level=3)

        rows = []
        for item in g_items:
            # 태스크 수 계산
            item_tasks = tasks_by_item.get(item["id"], [])
            task_count = len(item_tasks)
            hours = f"{item['estimated_hours']:.0f}h" if item["estimated_hours"] else "-"

            rows.append([
                item["code"] or "-",
                item["name"],
                item["description"] or "-",
                item["status"] or "-",
                hours,
                format_date(item["planned_start_date"]),
                format_date(item["planned_end_date"]),
            ])

            # 태스크가 있으면 하위 행 추가
            for task in item_tasks:
                t_hours = f"{task['estimated_hours']:.0f}h" if task["estimated_hours"] else "-"
                rows.append([
                    f"  └ {task['code'] or '-'}",
                    f"  {task['name']}",
                    task["description"] or "-",
                    task["status"] or "-",
                    t_hours,
                    format_date(task["planned_start_date"]),
                    format_date(task["planned_end_date"]),
                ])

        add_styled_table(doc,
            ["코드", "작업명", "설명", "상태", "공수", "시작일", "종료일"],
            rows,
            col_widths=[2, 3, 4, 1.8, 1.2, 2, 2],
        )
        doc.add_paragraph()


def add_section_5(doc, risks):
    """5. 리스크 관리"""
    doc.add_heading("5. 리스크 관리", level=1)
    doc.add_paragraph(
        "본 사업 수행 시 예상되는 주요 리스크와 대응 방안은 다음과 같습니다."
    )
    doc.add_paragraph()

    if risks:
        rows = []
        for risk in risks:
            rows.append([
                risk["title"],
                risk["category"] or "-",
                risk["impact"] or "-",
                risk["probability"] or "-",
                risk["description"] or "-",
                risk["mitigation_plan"] or "대응 계획 수립 필요",
            ])

        add_styled_table(doc,
            ["리스크명", "유형", "영향도", "발생확률", "설명", "대응 방안"],
            rows,
            col_widths=[2.5, 1.5, 1.5, 1.5, 4.5, 4.5],
        )
    else:
        doc.add_paragraph("(리스크 항목은 사업 착수 후 상세 식별 예정)")
    doc.add_paragraph()


def add_section_6(doc):
    """6. 제안 요구사항"""
    doc.add_heading("6. 제안 요구사항", level=1)

    doc.add_heading("6.1 제안서 구성", level=2)
    doc.add_paragraph("제안사는 다음 항목을 포함하여 제안서를 제출하여야 합니다:")
    doc.add_paragraph()

    proposal_items = [
        "사업 이해도 및 수행 방안",
        "시스템 아키텍처 설계",
        "기술 구현 방안 (요구사항별 대응 방안 포함)",
        "프로젝트 관리 방안 (WBS, 일정, 품질 관리)",
        "투입 인력 구성 및 역할",
        "유지보수 및 기술 지원 방안",
        "보안 대책",
        "사업비 산출 내역",
    ]
    for item in proposal_items:
        doc.add_paragraph(item, style="List Number")

    doc.add_paragraph()

    doc.add_heading("6.2 평가 기준", level=2)
    add_styled_table(doc,
        ["평가 항목", "배점", "평가 내용"],
        [
            ["기술 이해도", "20점", "요구사항 분석의 정확성, 과제 이해도"],
            ["기술 구현 방안", "30점", "아키텍처 설계, 기술 적용 방안의 적정성"],
            ["프로젝트 관리", "15점", "WBS, 일정, 품질 관리의 구체성"],
            ["투입 인력", "15점", "핵심 인력의 역량, 투입 인력의 적정성"],
            ["가격 적정성", "20점", "사업비 산출의 합리성"],
        ],
        col_widths=[4, 2, 10],
    )
    doc.add_paragraph()


def add_section_7(doc, project):
    """7. 계약 조건"""
    doc.add_heading("7. 계약 조건", level=1)

    items = [
        f"계약 기간: {format_date(project['start_date'])} ~ {format_date(project['end_date'])}",
        "계약 방식: 총액 도급 계약",
        "대금 지급: 착수금 30%, 중간금 30%, 잔금 40% (검수 완료 후)",
        "하자보수 기간: 검수 완료일로부터 1년",
        "지적재산권: 본 사업을 통해 생성된 모든 산출물의 소유권은 발주처에 귀속",
        "비밀유지: 사업 수행 중 취득한 정보에 대해 비밀유지 의무 부여",
        "보안: 개인정보보호법 및 정보보호 관련 법규 준수 필수",
    ]

    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()

    # 마무리
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 이상 —")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ── 메인 생성 함수 ───────────────────────────────────────────

def generate_rfp_docx(project, phases, groups, items, tasks, requirements, risks):
    """한 프로젝트에 대한 RFP Word 문서 생성"""
    doc = Document()

    # 문서 여백 설정
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 스타일 설정
    set_document_styles(doc)

    # 각 섹션 추가
    add_cover_page(doc, project)
    add_table_of_contents(doc)
    add_section_1(doc, project)
    add_section_2(doc, phases)
    add_section_3(doc, requirements)
    add_section_4(doc, phases, groups, items, tasks)
    add_section_5(doc, risks)
    add_section_6(doc)
    add_section_7(doc, project)

    return doc


def main():
    """메인 실행"""
    # 출력 디렉토리 생성
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    conn = get_connection()
    try:
        projects = fetch_projects(conn)
        print(f"총 {len(projects)}개 프로젝트 발견\n")

        for project in projects:
            print(f"프로젝트: {project['name']} ({project['id']})")

            # 데이터 조회
            phases = fetch_phases(conn, project["id"])
            groups = fetch_wbs_groups(conn, project["id"])
            items = fetch_wbs_items(conn, project["id"])
            tasks = fetch_wbs_tasks(conn, project["id"])
            requirements = fetch_requirements(conn, project["id"])
            risks = fetch_risks(conn, project["id"])

            print(f"  페이즈: {len(phases)}개, WBS 그룹: {len(groups)}개, "
                  f"아이템: {len(items)}개, 태스크: {len(tasks)}개")
            print(f"  요구사항: {len(requirements)}개, 리스크: {len(risks)}개")

            # 문서 생성
            doc = generate_rfp_docx(project, phases, groups, items, tasks, requirements, risks)

            # 파일명 생성 (프로젝트ID 기반)
            safe_name = project["name"].replace(" ", "_").replace("/", "_")
            filename = f"RFP_{project['id']}_{safe_name}.docx"
            filepath = os.path.join(OUTPUT_DIR, filename)

            doc.save(filepath)
            print(f"  → 저장 완료: {filepath}")
            print()

        print("=" * 60)
        print(f"총 {len(projects)}개 RFP 문서 생성 완료!")
        print(f"출력 경로: {os.path.abspath(OUTPUT_DIR)}")

    finally:
        conn.close()


if __name__ == "__main__":
    main()
